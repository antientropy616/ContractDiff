#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""测试文件保存"""

import os
from datetime import datetime
from docx import Document

# 测试路径
save_path = os.path.expanduser("~/Documents")
print(f"保存路径：{save_path}")
print(f"路径存在：{os.path.exists(save_path)}")
print(f"路径可写：{os.access(save_path, os.W_OK)}")

# 创建测试文档
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
output_file = os.path.join(save_path, f"合同对比结果_{timestamp}.docx")
print(f"输出文件：{output_file}")

try:
    doc = Document()
    doc.add_heading('测试文档', 0)
    doc.add_paragraph('这是一个测试文档，用于验证文件保存功能。')
    doc.save(output_file)
    print(f"✅ 文件保存成功：{output_file}")
    print(f"文件大小：{os.path.getsize(output_file)} bytes")
except Exception as e:
    print(f"❌ 保存失败：{e}")
    import traceback
    traceback.print_exc()
