#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
单元测试模块
测试合同对比助手的核心功能
"""

import unittest
import os
import sys
import tempfile
from datetime import datetime

# 添加 src 目录到路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from docx import Document


class TestDocumentComparator(unittest.TestCase):
    """测试文档对比器"""
    
    def setUp(self):
        """测试前准备"""
        # 创建测试文档
        self.temp_dir = tempfile.mkdtemp()
        
        # 创建测试文档 X
        self.doc_x_path = os.path.join(self.temp_dir, 'test_x.docx')
        doc_x = Document()
        doc_x.add_heading('测试文档', 0)
        doc_x.add_paragraph('这是第一段文字。')
        doc_x.add_paragraph('这是第二段文字，包含一些数字：123。')
        doc_x.add_paragraph('这是第三段文字，日期：2026-03-03。')
        doc_x.save(self.doc_x_path)
        
        # 创建测试文档 Y（有修改）
        self.doc_y_path = os.path.join(self.temp_dir, 'test_y.docx')
        doc_y = Document()
        doc_y.add_heading('测试文档', 0)
        doc_y.add_paragraph('这是第一段文字。')
        doc_y.add_paragraph('这是第二段文字，包含一些数字：456。')  # 数字修改
        doc_y.add_paragraph('这是新增的段落。')  # 新增段落
        doc_y.save(self.doc_y_path)
        
        # 导入对比器
        from core.doc_comparator import DocumentComparator
        self.comparator = DocumentComparator()
    
    def tearDown(self):
        """测试后清理"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def test_load_documents(self):
        """测试文档加载"""
        self.comparator.load_documents(self.doc_x_path, self.doc_y_path)
        
        self.assertIsNotNone(self.comparator.doc_x)
        self.assertIsNotNone(self.comparator.doc_y)
        self.assertEqual(len(self.comparator.paragraphs_x), 4)
        self.assertEqual(len(self.comparator.paragraphs_y), 4)
    
    def test_compare_content(self):
        """测试内容对比"""
        self.comparator.load_documents(self.doc_x_path, self.doc_y_path)
        self.comparator.analyze_structure()
        self.comparator.compare_content()
        
        # 应该有改动记录
        self.assertGreater(len(self.comparator.changes), 0)
    
    def test_identify_number_changes(self):
        """测试数字变化识别"""
        self.comparator.load_documents(self.doc_x_path, self.doc_y_path)
        self.comparator.analyze_structure()
        self.comparator.compare_content()
        
        # 应该识别到数字变化
        self.assertGreater(self.comparator.stats.get('number_changes', 0), 0)
    
    def test_generate_output(self):
        """测试输出生成"""
        self.comparator.load_documents(self.doc_x_path, self.doc_y_path)
        self.comparator.analyze_structure()
        self.comparator.compare_content()
        self.comparator.identify_changes()
        
        options = {
            'detail_list': True,
            'report': True,
            'save_path': self.temp_dir
        }
        
        output_path = self.comparator.generate_output(options)
        
        # 检查输出文件是否生成
        self.assertTrue(os.path.exists(output_path))
        self.assertTrue(output_path.endswith('.docx'))


class TestVersionComparison(unittest.TestCase):
    """测试版本号比较"""
    
    def test_version_compare(self):
        """测试版本号比较逻辑"""
        from utils.updater import AutoUpdater
        
        updater = AutoUpdater()
        
        # 测试大于
        self.assertEqual(updater._compare_versions('1.2', '1.1'), 1)
        self.assertEqual(updater._compare_versions('2.0', '1.9'), 1)
        
        # 测试等于
        self.assertEqual(updater._compare_versions('1.2', '1.2'), 0)
        self.assertEqual(updater._compare_versions('1.0', '1.0'), 0)
        
        # 测试小于
        self.assertEqual(updater._compare_versions('1.1', '1.2'), -1)
        self.assertEqual(updater._compare_versions('1.9', '2.0'), -1)


class TestLogger(unittest.TestCase):
    """测试日志系统"""
    
    def test_logger_initialization(self):
        """测试日志器初始化"""
        from utils.logger import Logger
        
        logger = Logger()
        self.assertIsNotNone(logger.logger)
    
    def test_logger_methods(self):
        """测试日志方法"""
        from utils.logger import Logger
        
        logger = Logger()
        
        # 测试各日志级别不抛出异常
        logger.debug("测试调试日志")
        logger.info("测试信息日志")
        logger.warning("测试警告日志")
        logger.error("测试错误日志")


class TestErrorHandler(unittest.TestCase):
    """测试错误处理"""
    
    def test_custom_exceptions(self):
        """测试自定义异常"""
        from utils.error_handler import (
            ContractDiffError,
            FileLoadError,
            FileFormatError,
            CompareError,
            OutputError
        )
        
        # 测试基础异常
        try:
            raise ContractDiffError("测试错误", "TEST_ERROR")
        except ContractDiffError as e:
            self.assertEqual(e.error_code, "TEST_ERROR")
        
        # 测试文件加载错误
        try:
            raise FileLoadError("无法加载", "/path/to/file.docx")
        except FileLoadError as e:
            self.assertIn("/path/to/file.docx", str(e))
        
        # 测试文件格式错误
        try:
            raise FileFormatError("格式错误", ".docx")
        except FileFormatError as e:
            self.assertIn(".docx", str(e))
    
    def test_friendly_error_message(self):
        """测试友好错误消息"""
        from utils.error_handler import get_friendly_error_message, FileLoadError
        
        error = FileLoadError("测试错误")
        message = get_friendly_error_message(error)
        
        self.assertIsInstance(message, str)
        self.assertGreater(len(message), 0)


class TestPDFProcessor(unittest.TestCase):
    """测试 PDF 处理器"""
    
    def test_pdf_format_check(self):
        """测试 PDF 格式检查"""
        from utils.pdf_handler import PDFProcessor
        
        processor = PDFProcessor()
        
        self.assertTrue(processor.is_pdf('test.pdf'))
        self.assertTrue(processor.is_pdf('test.PDF'))
        self.assertFalse(processor.is_pdf('test.docx'))
        self.assertFalse(processor.is_pdf('test.txt'))


def run_tests():
    """运行所有测试"""
    # 创建测试套件
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()
    
    # 添加测试
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentComparator))
    suite.addTests(loader.loadTestsFromTestCase(TestVersionComparison))
    suite.addTests(loader.loadTestsFromTestCase(TestLogger))
    suite.addTests(loader.loadTestsFromTestCase(TestErrorHandler))
    suite.addTests(loader.loadTestsFromTestCase(TestPDFProcessor))
    
    # 运行测试
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # 返回测试结果
    return result.wasSuccessful()


if __name__ == '__main__':
    success = run_tests()
    sys.exit(0 if success else 1)
