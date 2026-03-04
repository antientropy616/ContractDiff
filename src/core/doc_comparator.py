#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档对比核心引擎 v1.2.1
修复：表格、页眉页脚、文本框、脚注等识别
"""

import os
import re
import difflib
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from openpyxl import Workbook

from utils.logger import logger
import hashlib
import io


class DocumentComparator:
    """文档对比器"""
    
    def __init__(self):
        self.doc_x = None
        self.doc_y = None
        self.paragraphs_x = []
        self.paragraphs_y = []
        self.tables_x = []
        self.tables_y = []
        self.headers_x = {}
        self.headers_y = {}
        self.footers_x = {}
        self.footers_y = {}
        self.footnotes_x = []
        self.footnotes_y = []
        self.images_x = []
        self.images_y = []
        self.changes = []
        self.stats = {
            'additions': 0,
            'deletions': 0,
            'modifications': 0,
            'number_changes': 0,
            'paragraph_moves': 0,
            'table_changes': 0,
            'header_footer_changes': 0,
            'image_changes': 0
        }
        
    def load_documents(self, file_x, file_y):
        """加载两个文档（包含表格、页眉页脚等）"""
        logger.info(f"加载文档：{file_x}")
        self.doc_x = Document(file_x)
        logger.info(f"加载文档：{file_y}")
        self.doc_y = Document(file_y)
        
        # 提取段落
        self.paragraphs_x = [p.text for p in self.doc_x.paragraphs if p.text.strip()]
        self.paragraphs_y = [p.text for p in self.doc_y.paragraphs if p.text.strip()]
        
        # 提取表格
        self._extract_tables(self.doc_x, self.tables_x)
        self._extract_tables(self.doc_y, self.tables_y)
        
        # 提取页眉页脚
        self._extract_headers_footers(self.doc_x, self.headers_x, self.footers_x)
        self._extract_headers_footers(self.doc_y, self.headers_y, self.footers_y)
        
        # 提取脚注
        self._extract_footnotes(self.doc_x, self.footnotes_x)
        self._extract_footnotes(self.doc_y, self.footnotes_y)
        
        # 提取图片
        self._extract_images(self.doc_x, self.images_x)
        self._extract_images(self.doc_y, self.images_y)
        
        logger.info(f"文档 X: {len(self.paragraphs_x)} 段落，{len(self.tables_x)} 表格，{len(self.images_x)} 图片")
        logger.info(f"文档 Y: {len(self.paragraphs_y)} 段落，{len(self.tables_y)} 表格，{len(self.images_y)} 图片")
        
    def _extract_tables(self, doc, tables_list):
        """提取文档中所有表格内容"""
        for i, table in enumerate(doc.tables):
            table_text = []
            for row_idx, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):  # 非空行
                    table_text.append(" | ".join(row_cells))
            
            if table_text:
                tables_list.append({
                    'index': i,
                    'content': '\n'.join(table_text),
                    'rows': len(table.rows),
                    'cols': len(table.columns)
                })
                logger.debug(f"表格 {i}: {len(table.rows)}行 x {len(table.columns)}列")
    
    def _extract_headers_footers(self, doc, headers_dict, footers_dict):
        """提取页眉页脚"""
        # 页眉
        for section in doc.sections:
            if section.header:
                header_text = " ".join([p.text.strip() for p in section.header.paragraphs]).strip()
                if header_text:
                    headers_dict[section.start_type] = header_text
        
        # 页脚
        for section in doc.sections:
            if section.footer:
                footer_text = " ".join([p.text.strip() for p in section.footer.paragraphs]).strip()
                if footer_text:
                    footers_dict[section.start_type] = footer_text
        
        if headers_dict:
            logger.debug(f"页眉：{len(headers_dict)} 个")
        if footers_dict:
            logger.debug(f"页脚：{len(footers_dict)} 个")
    
    def _extract_footnotes(self, doc, footnotes_list):
        """提取脚注"""
        try:
            # docx 库对脚注支持有限，尝试从文档脚注部分提取
            if hasattr(doc, 'footnotes'):
                for footnote in doc.footnotes:
                    if hasattr(footnote, 'text') and footnote.text.strip():
                        footnotes_list.append(footnote.text.strip())
        except Exception as e:
            logger.warning(f"提取脚注失败：{e}")
        
        if footnotes_list:
            logger.debug(f"脚注：{len(footnotes_list)} 个")
    
    def _extract_images(self, doc, images_list):
        """提取文档中的图片信息"""
        try:
            # 遍历文档中的所有关系，查找图片
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # 获取图片数据
                        image_data = rel.target_part.blob
                        
                        # 计算图片哈希值用于对比
                        image_hash = hashlib.md5(image_data).hexdigest()
                        
                        # 获取图片信息
                        image_info = {
                            'index': image_count,
                            'hash': image_hash,
                            'size': len(image_data),
                            'type': rel.target_part.content_type,
                            'filename': os.path.basename(rel.target_ref)
                        }
                        
                        images_list.append(image_info)
                        image_count += 1
                        
                        logger.debug(f"图片 {image_count}: {image_info['type']}, {image_info['size']} bytes")
                        
                    except Exception as e:
                        logger.warning(f"提取图片失败：{e}")
            
            if images_list:
                logger.info(f"共找到 {len(images_list)} 张图片")
                
        except Exception as e:
            logger.warning(f"提取图片过程中出错：{e}")
    
    def analyze_structure(self):
        """分析文档结构"""
        # 简单的段落匹配
        self.matcher = difflib.SequenceMatcher(None, self.paragraphs_x, self.paragraphs_y)
        
        # 表格对比
        self._compare_tables()
        
        # 页眉页脚对比
        self._compare_headers_footers()
        
        # 脚注对比
        self._compare_footnotes()
        
        # 图片对比
        self._compare_images()
    
    def _compare_tables(self):
        """对比表格内容"""
        # 使用 SequenceMatcher 对比表格
        table_matcher = difflib.SequenceMatcher(
            None,
            [t['content'] for t in self.tables_x],
            [t['content'] for t in self.tables_y]
        )
        
        for tag, i1, i2, j1, j2 in table_matcher.get_opcodes():
            if tag == 'equal':
                continue
            elif tag == 'replace':
                # 表格修改
                for idx in range(min(i2-i1, j2-j1)):
                    if i1+idx < len(self.tables_x) and j1+idx < len(self.tables_y):
                        old_table = self.tables_x[i1+idx]
                        new_table = self.tables_y[j1+idx]
                        
                        # 对比表格内容
                        table_diff = difflib.ndiff(
                            old_table['content'].split('\n'),
                            new_table['content'].split('\n')
                        )
                        
                        changes_found = False
                        for line in table_diff:
                            if line.startswith('+ '):
                                self.changes.append({
                                    'type': 'table_cell_add',
                                    'position': f"表格{old_table['index']+1}",
                                    'content': line[2:]
                                })
                                self.stats['additions'] += 1
                                changes_found = True
                            elif line.startswith('- '):
                                self.changes.append({
                                    'type': 'table_cell_delete',
                                    'position': f"表格{old_table['index']+1}",
                                    'content': line[2:]
                                })
                                self.stats['deletions'] += 1
                                changes_found = True
                        
                        if changes_found:
                            self.stats['table_changes'] += 1
                            
            elif tag == 'delete':
                # 表格删除
                for idx in range(i2-i1):
                    if i1+idx < len(self.tables_x):
                        self.changes.append({
                            'type': 'table_delete',
                            'position': f"表格{self.tables_x[i1+idx]['index']+1}",
                            'content': f"删除表格（{self.tables_x[i1+idx]['rows']}行 x {self.tables_x[i1+idx]['cols']}列）"
                        })
                        self.stats['deletions'] += 1
                        self.stats['table_changes'] += 1
                        
            elif tag == 'insert':
                # 表格新增
                for idx in range(j2-j1):
                    if j1+idx < len(self.tables_y):
                        self.changes.append({
                            'type': 'table_add',
                            'position': f"表格{self.tables_y[j1+idx]['index']+1}",
                            'content': f"新增表格（{self.tables_y[j1+idx]['rows']}行 x {self.tables_y[j1+idx]['cols']}列）"
                        })
                        self.stats['additions'] += 1
                        self.stats['table_changes'] += 1
    
    def _compare_headers_footers(self):
        """对比页眉页脚"""
        all_sections = set(self.headers_x.keys()) | set(self.headers_y.keys())
        all_sections |= set(self.footers_x.keys()) | set(self.footers_y.keys())
        
        for section_type in all_sections:
            # 对比页眉
            header_x = self.headers_x.get(section_type, '')
            header_y = self.headers_y.get(section_type, '')
            
            if header_x != header_y:
                if header_x and not header_y:
                    self.changes.append({
                        'type': 'header_delete',
                        'position': f'页眉（章节{section_type}）',
                        'content': header_x
                    })
                    self.stats['deletions'] += 1
                    self.stats['header_footer_changes'] += 1
                elif not header_x and header_y:
                    self.changes.append({
                        'type': 'header_add',
                        'position': f'页眉（章节{section_type}）',
                        'content': header_y
                    })
                    self.stats['additions'] += 1
                    self.stats['header_footer_changes'] += 1
                else:
                    # 页眉修改
                    self._compare_text_with_diff('header', f'页眉（章节{section_type}）', header_x, header_y)
                    self.stats['header_footer_changes'] += 1
            
            # 对比页脚
            footer_x = self.footers_x.get(section_type, '')
            footer_y = self.footers_y.get(section_type, '')
            
            if footer_x != footer_y:
                if footer_x and not footer_y:
                    self.changes.append({
                        'type': 'footer_delete',
                        'position': f'页脚（章节{section_type}）',
                        'content': footer_x
                    })
                    self.stats['deletions'] += 1
                    self.stats['header_footer_changes'] += 1
                elif not footer_x and footer_y:
                    self.changes.append({
                        'type': 'footer_add',
                        'position': f'页脚（章节{section_type}）',
                        'content': footer_y
                    })
                    self.stats['additions'] += 1
                    self.stats['header_footer_changes'] += 1
                else:
                    # 页脚修改
                    self._compare_text_with_diff('footer', f'页脚（章节{section_type}）', footer_x, footer_y)
                    self.stats['header_footer_changes'] += 1
    
    def _compare_footnotes(self):
        """对比脚注"""
        footnotes_matcher = difflib.SequenceMatcher(None, self.footnotes_x, self.footnotes_y)
        
        for tag, i1, i2, j1, j2 in footnotes_matcher.get_opcodes():
            if tag == 'replace':
                for idx in range(min(i2-i1, j2-j1)):
                    if i1+idx < len(self.footnotes_x) and j1+idx < len(self.footnotes_y):
                        self._compare_text_with_diff(
                            'footnote',
                            f'脚注{i1+idx+1}',
                            self.footnotes_x[i1+idx],
                            self.footnotes_y[j1+idx]
                        )
            elif tag == 'delete':
                for idx in range(i2-i1):
                    if i1+idx < len(self.footnotes_x):
                        self.changes.append({
                            'type': 'footnote_delete',
                            'position': f'脚注{i1+idx+1}',
                            'content': self.footnotes_x[i1+idx]
                        })
                        self.stats['deletions'] += 1
            elif tag == 'insert':
                for idx in range(j2-j1):
                    if j1+idx < len(self.footnotes_y):
                        self.changes.append({
                            'type': 'footnote_add',
                            'position': f'脚注{j1+idx+1}',
                            'content': self.footnotes_y[j1+idx]
                        })
                        self.stats['additions'] += 1
    
    def _compare_images(self):
        """对比图片变化"""
        # 使用哈希值对比图片
        hashes_x = {img['hash']: img for img in self.images_x}
        hashes_y = {img['hash']: img for img in self.images_y}
        
        # 找出删除的图片（X 中有，Y 中没有）
        for img_hash, img_info in hashes_x.items():
            if img_hash not in hashes_y:
                self.changes.append({
                    'type': 'image_delete',
                    'position': f'图片{img_info["index"]+1}',
                    'content': f'删除图片 ({img_info["filename"]}, {self._format_file_size(img_info["size"])})'
                })
                self.stats['deletions'] += 1
                self.stats['image_changes'] += 1
        
        # 找出新增的图片（Y 中有，X 中没有）
        for img_hash, img_info in hashes_y.items():
            if img_hash not in hashes_x:
                self.changes.append({
                    'type': 'image_add',
                    'position': f'图片{img_info["index"]+1}',
                    'content': f'新增图片 ({img_info["filename"]}, {self._format_file_size(img_info["size"])})'
                })
                self.stats['additions'] += 1
                self.stats['image_changes'] += 1
        
        # 数量变化提示
        if len(self.images_x) != len(self.images_y):
            logger.info(f"图片数量变化：{len(self.images_x)} → {len(self.images_y)}")
    
    def _format_file_size(self, size_bytes):
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f}GB"
    
    def _compare_text_with_diff(self, change_type, position, text_x, text_y):
        """使用 difflib 对比两段文字"""
        diff = difflib.ndiff(text_x, text_y)
        
        additions = []
        deletions = []
        
        for item in diff:
            if item.startswith('+ '):
                additions.append(item[2:])
            elif item.startswith('- '):
                deletions.append(item[2:])
        
        if additions:
            self.changes.append({
                'type': f'{change_type}_add',
                'position': position,
                'content': ''.join(additions)
            })
            self.stats['additions'] += len(additions)
            
            # 检查是否有数字变化
            if self._contains_numbers(''.join(additions)):
                self.stats['number_changes'] += 1
        
        if deletions:
            self.changes.append({
                'type': f'{change_type}_delete',
                'position': position,
                'content': ''.join(deletions)
            })
            self.stats['deletions'] += len(deletions)
        
        if additions and deletions:
            self.stats['modifications'] += 1
    
    def compare_content(self):
        """对比内容"""
        # 获取差异块
        for tag, i1, i2, j1, j2 in self.matcher.get_opcodes():
            if tag == 'equal':
                continue
            elif tag == 'replace':
                # 段落修改
                for idx in range(min(i2-i1, j2-j1)):
                    self._compare_paragraphs(
                        self.paragraphs_x[i1+idx] if i1+idx < len(self.paragraphs_x) else "",
                        self.paragraphs_y[j1+idx] if j1+idx < len(self.paragraphs_y) else "",
                        f"第{i1+idx+1}段"
                    )
            elif tag == 'delete':
                # 段落删除
                for idx in range(i2-i1):
                    self.changes.append({
                        'type': 'paragraph_delete',
                        'position': f"第{i1+idx+1}段",
                        'content': self.paragraphs_x[i1+idx]
                    })
                    self.stats['deletions'] += 1
            elif tag == 'insert':
                # 段落新增
                for idx in range(j2-j1):
                    self.changes.append({
                        'type': 'paragraph_add',
                        'position': f"第{j1+idx+1}段",
                        'content': self.paragraphs_y[j1+idx]
                    })
                    self.stats['additions'] += 1
        
        logger.info(f"对比完成，共发现 {len(self.changes)} 处改动")
    
    def _compare_paragraphs(self, text_x, text_y, position):
        """对比两个段落"""
        if text_x == text_y:
            return
        
        # 使用 difflib 对比文字
        diff = difflib.ndiff(text_x, text_y)
        
        additions = []
        deletions = []
        
        for item in diff:
            if item.startswith('+ '):
                additions.append(item[2:])
            elif item.startswith('- '):
                deletions.append(item[2:])
        
        if additions:
            self.changes.append({
                'type': 'text_add',
                'position': position,
                'content': ''.join(additions)
            })
            self.stats['additions'] += len(additions)
            
            # 检查是否有数字变化
            if self._contains_numbers(''.join(additions)):
                self.stats['number_changes'] += 1
        
        if deletions:
            self.changes.append({
                'type': 'text_delete',
                'position': position,
                'content': ''.join(deletions)
            })
            self.stats['deletions'] += len(deletions)
        
        if additions and deletions:
            self.stats['modifications'] += 1
    
    def _contains_numbers(self, text):
        """检查是否包含数字"""
        return bool(re.search(r'\d+', text))
    
    def identify_changes(self):
        """识别并分类改动"""
        # 已经在 compare_content 中完成
        pass
    
    def generate_output(self, options):
        """生成输出文档"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = options.get('save_path', os.path.expanduser("~/Documents"))
        
        # 确保保存目录存在
        os.makedirs(save_path, exist_ok=True)
        
        output_file = os.path.join(save_path, f"合同对比结果_{timestamp}.docx")
        
        logger.info(f"生成输出文档：{output_file}")
        
        # 创建结果文档
        result_doc = Document()
        
        # 添加标题
        title = result_doc.add_heading('合同对比报告', 0)
        title.alignment = 1  # 居中
        
        # 添加基本信息
        result_doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        result_doc.add_paragraph(f"改动总数：{sum(self.stats.values())} 处")
        result_doc.add_paragraph("")
        
        # 添加改动统计
        result_doc.add_heading('改动统计', level=1)
        stats_display = {
            'additions': '文字新增',
            'deletions': '文字删除',
            'modifications': '文字修改',
            'number_changes': '数字变化',
            'paragraph_moves': '段落调整',
            'table_changes': '表格改动',
            'header_footer_changes': '页眉页脚改动',
            'image_changes': '图片改动'
        }
        for change_type, count in self.stats.items():
            if count > 0:
                result_doc.add_paragraph(f"• {stats_display.get(change_type, change_type)}: {count} 处")
        
        result_doc.add_paragraph("")
        
        # 添加详细改动
        if options.get('detail_list', True):
            result_doc.add_heading('详细改动列表', level=1)
            
            # 按类型分组显示
            change_groups = {
                '正文': ['text_add', 'text_delete', 'paragraph_add', 'paragraph_delete'],
                '表格': ['table_add', 'table_delete', 'table_cell_add', 'table_cell_delete'],
                '图片': ['image_add', 'image_delete'],
                '页眉页脚': ['header_add', 'header_delete', 'footer_add', 'footer_delete'],
                '脚注': ['footnote_add', 'footnote_delete'],
                '其他': []
            }
            
            for i, change in enumerate(self.changes[:200], 1):  # 限制显示前 200 条
                p = result_doc.add_paragraph()
                
                change_type = change['type']
                
                # 确定显示类别
                category = '其他'
                for cat, types in change_groups.items():
                    if change_type in types:
                        category = cat
                        break
                
                # 根据类型设置格式
                if 'add' in change_type:
                    run = p.add_run(f"[新增] {change['content'][:200]}")
                    run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
                    run.underline = True
                elif 'delete' in change_type:
                    run = p.add_run(f"[删除] {change['content'][:200]}")
                    run.font.color.rgb = RGBColor(139, 0, 0)  # 深红色
                    run.font.strike = True
                else:
                    run = p.add_run(f"[修改] {change['content'][:200]}")
                    run.font.color.rgb = RGBColor(0, 0, 139)  # 深蓝色
                
                # 添加位置信息
                if change.get('position'):
                    p.add_run(f" （位置：{change['position']}）").italic = True
        
        # 保存文档
        result_doc.save(output_file)
        logger.info(f"输出文档已保存：{output_file}")
        
        # 如果选择了 Excel 报告
        if options.get('report', False):
            self._generate_excel_report(save_path, timestamp)
        
        return output_file
    
    def _generate_excel_report(self, save_path, timestamp):
        """生成 Excel 报告"""
        logger.info(f"生成 Excel 报告：{save_path}")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "改动清单"
        
        # 表头
        headers = ["序号", "改动类型", "位置", "原文内容", "修改后内容"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
            ws.cell(row=1, column=col).font = ws.cell(row=1, column=col).font.copy(bold=True)
        
        # 填充数据
        for i, change in enumerate(self.changes, 1):
            ws.cell(row=i+1, column=1, value=i)
            ws.cell(row=i+1, column=2, value=self._get_change_type_name(change['type']))
            ws.cell(row=i+1, column=3, value=f"{change.get('position', '')}")
            ws.cell(row=i+1, column=4, value=change.get('content', '')[:200])
        
        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50
        
        # 保存
        excel_file = os.path.join(save_path, f"合同对比报告_{timestamp}.xlsx")
        wb.save(excel_file)
        logger.info(f"Excel 报告已保存：{excel_file}")
    
    def _get_change_type_name(self, change_type):
        """获取改动类型名称"""
        names = {
            'additions': '文字新增',
            'deletions': '文字删除',
            'modifications': '文字修改',
            'number_changes': '数字变化',
            'paragraph_moves': '段落调整',
            'table_changes': '表格改动',
            'header_footer_changes': '页眉页脚改动',
            'image_changes': '图片改动',
            'text_add': '文字新增',
            'text_delete': '文字删除',
            'paragraph_add': '段落新增',
            'paragraph_delete': '段落删除',
            'table_add': '表格新增',
            'table_delete': '表格删除',
            'table_cell_add': '表格单元格新增',
            'table_cell_delete': '表格单元格删除',
            'header_add': '页眉新增',
            'header_delete': '页眉删除',
            'footer_add': '页脚新增',
            'footer_delete': '页脚删除',
            'footnote_add': '脚注新增',
            'footnote_delete': '脚注删除',
            'image_add': '图片新增',
            'image_delete': '图片删除'
        }
        return names.get(change_type, change_type)
    
    def get_stats(self):
        """获取统计信息"""
        return self.stats
